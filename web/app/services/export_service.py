import os
from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from reportlab.lib.pagesizes import A4
import re

from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from flask import current_app
from web.app.models import DocumentApp, DocumentSectionContent
from docx.shared import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls


def clean_html_content(html_content):
    """
    Clean HTML content by removing tags and preserving text formatting.
    
    Args:
        html_content (str): HTML content to clean
        
    Returns:
        str: Cleaned text content
    """
    if not html_content:
        return ""
        
    # Replace <br> and <p> tags with newlines
    content = re.sub(r'<br\s*/?>', '\n', html_content, flags=re.IGNORECASE)
    content = re.sub(r'</?p\s*/?>', '\n', content, flags=re.IGNORECASE)
    
    # Replace <i> and <em> tags with special markers
    content = re.sub(r'<i>', '[[italic_start]]', content, flags=re.IGNORECASE)
    content = re.sub(r'</i>', '[[italic_end]]', content, flags=re.IGNORECASE)
    content = re.sub(r'<em>', '[[italic_start]]', content, flags=re.IGNORECASE)
    content = re.sub(r'</em>', '[[italic_end]]', content, flags=re.IGNORECASE)
    
    # Replace <h1> to <h6> tags with special markers
    for i in range(1, 7):
        content = re.sub(fr'<h{i}>', f'[[heading_start_{i}]]', content, flags=re.IGNORECASE)
        content = re.sub(fr'</h{i}>', f'[[heading_end_{i}]]', content, flags=re.IGNORECASE)
    
    # Preserve list formatting
    content = re.sub(r'<ol>', '[[list_start]]', content, flags=re.IGNORECASE)
    content = re.sub(r'</ol>', '[[list_end]]', content, flags=re.IGNORECASE)
    content = re.sub(r'<li>', '[[list_item]]', content, flags=re.IGNORECASE)
    content = re.sub(r'</li>', '[[list_item_end]]', content, flags=re.IGNORECASE)
    
    # Remove all other HTML tags
    content = re.sub(r'<[^>]+>', '', content)
    
    # Replace HTML entities
    content = content.replace('&nbsp;', ' ')
    content = content.replace('&amp;', '&')
    content = content.replace('&lt;', '<')
    content = content.replace('&gt;', '>')
    content = content.replace('&quot;', '"')
    content = content.replace('&#39;', "'")
    
    # Clean up multiple newlines and whitespace
    content = re.sub(r'\n\s*\n', '\n', content)  # Replace multiple newlines with single newline
    content = re.sub(r' +', ' ', content)
    
    # Remove leading/trailing whitespace and empty lines
    content = '\n'.join(line for line in content.split('\n') if line.strip())
    content = content.strip()
    
    return content


# Constants for document formatting
DOCUMENT_SETTINGS = {
    'page_height': Cm(29.7),  # A4
    'page_width': Cm(21.0),   # A4
    'left_margin': Cm(3.0),   # Left margin 3 cm
    'right_margin': Cm(1.5),  # Right margin 1.5 cm
    'top_margin': Cm(2.0),    # Top margin 2 cm
    'bottom_margin': Cm(2.0), # Bottom margin 2 cm
    'title_font_size': Pt(14),
    'body_font_size': Pt(14),
    'first_line_indent': Cm(1.5),
    'line_spacing': WD_LINE_SPACING.ONE_POINT_FIVE
}

def setup_document_formatting(docx):
    """Set up document formatting according to GOST 7.32-2017"""
    section = docx.sections[0]
    section.page_height = DOCUMENT_SETTINGS['page_height']
    section.page_width = DOCUMENT_SETTINGS['page_width']
    section.left_margin = DOCUMENT_SETTINGS['left_margin']
    section.right_margin = DOCUMENT_SETTINGS['right_margin']
    section.top_margin = DOCUMENT_SETTINGS['top_margin']
    section.bottom_margin = DOCUMENT_SETTINGS['bottom_margin']

def create_paragraph(docx, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None):
    """Create a new paragraph with standard formatting"""
    p = docx.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing_rule = DOCUMENT_SETTINGS['line_spacing']
    p.paragraph_format.first_line_indent = first_line_indent if first_line_indent is not None else DOCUMENT_SETTINGS['first_line_indent']
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = False
    return p

def add_text_run(p, text, is_bold=False, is_italic=False):
    """Add a text run to a paragraph with standard formatting"""
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = DOCUMENT_SETTINGS['body_font_size']
    run.bold = is_bold
    run.italic = is_italic
    return run

def process_italic_text(p, text):
    """Process text with italic formatting markers"""
    parts = text.split('[[italic_start]]')
    
    # Add first part if it exists
    if parts[0]:
        is_bold = parts[0].strip().startswith('**') or parts[0].strip().startswith('__')
        if is_bold:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_run(p, parts[0], is_bold=is_bold)
    
    # Process remaining parts with italic formatting
    for part in parts[1:]:
        if '[[italic_end]]' in part:
            italic_part, rest = part.split('[[italic_end]]', 1)
            # Add italic part
            add_text_run(p, italic_part, is_italic=True)
            # Add rest of the text
            if rest:
                is_bold = rest.strip().startswith('**') or rest.strip().startswith('__')
                if is_bold:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_text_run(p, rest, is_bold=is_bold)
        else:
            # If no end marker, treat as normal text
            is_bold = part.strip().startswith('**') or part.strip().startswith('__')
            if is_bold:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text_run(p, part, is_bold=is_bold)

def process_list_items(docx, list_items):
    """Process list items and add them to the document"""
    for i, item in enumerate(list_items[1:], 1):  # Skip first empty item
        item = item.split('[[list_item_end]]')[0].strip()
        if item:
            p = create_paragraph(docx, first_line_indent=Cm(1.25))
            
            # Add list number
            add_text_run(p, f"{i}. ")
            
            # Add list item text
            process_italic_text(p, item)

def add_page_number(paragraph):
    """Add page number to paragraph"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
    run._element.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText.text = "PAGE"
    run._element.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
    run._element.append(fldChar2)

def add_ToC(paragraph):
    # Add a new section for the table of contents
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Нажмите правой кнопкой мыши и выберите 'Обновить поле'"
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)


def process_heading(docx, heading_text, level):
    """Process and add a heading to the document with appropriate formatting"""
    alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p = create_paragraph(docx, alignment=alignment, first_line_indent=Cm(0))
    add_text_run(p, heading_text, is_bold=True)
    
    # Set heading level using XML
    p.style = docx.styles['Normal']  # Use normal style to avoid Word styles
    outline_lvl = OxmlElement('w:outlineLvl')
    outline_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(level - 1))
    p._element.get_or_add_pPr().append(outline_lvl)

def export_to_docx(document_id):
    """
    Export document to DOCX format according to GOST 7.32-2017

    Args:
        document_id (int): Document ID

    Returns:
        str: Path to the exported file
    """
    document = DocumentApp.query.get_or_404(document_id)
    sections = document.sections.order_by(DocumentSectionContent.id).all()

    # Create a new Word document
    docx = DocxDocument()
    setup_document_formatting(docx)

    # Add page numbers to footer
    section = docx.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(paragraph)

    # Add title (assuming title is on the first page)
    title_paragraph = create_paragraph(docx, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text_run(title_paragraph, document.title, is_bold=True)
    title_paragraph.runs[0].font.size = DOCUMENT_SETTINGS['title_font_size']

    # Add a new page for the Table of Contents
    docx.add_section(WD_SECTION.NEW_PAGE)
    
    # Add the "Оглавление" heading
    toc_heading = create_paragraph(docx, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Cm(0))
    add_text_run(toc_heading, "Оглавление", is_bold=True)
    toc_heading.runs[0].font.size = DOCUMENT_SETTINGS['title_font_size'] # Match title size
    
    # Add the Table of Content field
    toc_field_paragraph = docx.add_paragraph()
    add_ToC(toc_field_paragraph)
    
    # Add document sections (starting on a new page after ToC)
    docx.add_section(WD_SECTION.NEW_PAGE) # Add new page break before main content
    for section_content in sections:
        content = clean_html_content(section_content.content)
        if not content.strip():  # Skip empty sections
            continue

        # Process content
        paragraphs = content.split('\n')
        for para in paragraphs:
            if not para.strip():
                continue

            # Check for headings
            for i in range(1, 7):
                if f'[[heading_start_{i}]]' in para:
                    heading_text = para.split(f'[[heading_start_{i}]]')[1].split(f'[[heading_end_{i}]]')[0].strip()
                    process_heading(docx, heading_text, level=i)
                    para = para.split(f'[[heading_end_{i}]]')[1] if f'[[heading_end_{i}]]' in para else ''

            # Check if paragraph contains a list
            if '[[list_start]]' in para:
                # Process content before list
                before_list = para.split('[[list_start]]')[0].strip()
                if before_list:
                    p = create_paragraph(docx)
                    process_italic_text(p, before_list)

                # Process list
                list_items = para.split('[[list_start]]')[1].split('[[list_end]]')[0].split('[[list_item]]')
                process_list_items(docx, list_items)

                # Process content after list
                after_list = para.split('[[list_end]]')[1].strip() if '[[list_end]]' in para else ''
                if after_list:
                    p = create_paragraph(docx)
                    process_italic_text(p, after_list)
            else:
                # Process regular paragraph
                p = create_paragraph(docx)
                process_italic_text(p, para)

    # Save the document
    output_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'exports')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{document.title}_{document.id}.docx")
    docx.save(output_path)

    return output_path


def export_to_pdf(document_id):
    """
    Export document to PDF format according to GOST 7.32-2017

    Args:
        document_id (int): Document ID

    Returns:
        str: Path to the exported file
    """
    document = DocumentApp.query.get_or_404(document_id)
    sections = document.sections.order_by(DocumentSectionContent.id).all()

    # Create output directory
    output_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'exports')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{document.title}_{document.id}.pdf")

    # Create PDF document
    pdf = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=85,  # ~3 cm
        rightMargin=42,  # ~1.5 cm
        topMargin=57,  # ~2 cm
        bottomMargin=57  # ~2 cm
    )

    # Define styles
    styles = getSampleStyleSheet()

    # Create custom styles according to GOST 7.32-2017
    gost_title_style = ParagraphStyle(
        name='GOSTTitle',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=16,
        leading=18,
        alignment=1,  # Center
        spaceAfter=12,
        encoding='UTF-8'
    )
    styles.add(gost_title_style)

    gost_heading_style = ParagraphStyle(
        name='GOSTHeading',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=14,
        leading=16,
        alignment=1,  # Center
        spaceAfter=12,
        encoding='UTF-8'
    )
    styles.add(gost_heading_style)

    gost_normal_style = ParagraphStyle(
        name='GOSTNormal',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=14,
        leading=21,  # 1.5 line spacing
        alignment=4,  # Justify
        firstLineIndent=35,  # 1.25 cm
        spaceBefore=0,
        spaceAfter=0,
        encoding='UTF-8'
    )
    styles.add(gost_normal_style)

    gost_bold_style = ParagraphStyle(
        name='GOSTBold',
        parent=styles['Normal'],
        fontName='Times-Bold',
        fontSize=14,
        leading=21,  # 1.5 line spacing
        alignment=1,  # Center
        spaceBefore=0,
        spaceAfter=0,
        encoding='UTF-8'
    )
    styles.add(gost_bold_style)

    # Build the document
    elements = []

    # Add title
    title_text = document.title.encode('utf-8').decode('utf-8')
    elements.append(Paragraph(title_text, styles['GOSTTitle']))
    elements.append(Spacer(1, 12))

    # Add sections
    for section_content in sections:
        section_template = section_content.section_template

        if section_template.code == 'ТЛ' or section_template.code == 'СИ':
        # Add section title
            title_text = section_template.name.encode('utf-8').decode('utf-8')
            elements.append(Paragraph(title_text, styles['GOSTHeading']))
            elements.append(Spacer(1, 12))

        # Add section content
        content = clean_html_content(section_content.content)
        if content.strip():
            # Split content into paragraphs and add each one
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    # Process italic formatting
                    parts = para.split('[[italic_start]]')
                    if len(parts) > 1:
                        # Handle italic text
                        for i, part in enumerate(parts):
                            if i == 0:
                                # First part is normal text
                                if part:
                                    text = part.encode('utf-8').decode('utf-8')
                                    if text.strip().startswith('**') or text.strip().startswith('__'):
                                        text = text.strip('*_').strip()
                                        elements.append(Paragraph(text, styles['GOSTBold']))
                                    else:
                                        elements.append(Paragraph(text, styles['GOSTNormal']))
                            else:
                                # Handle italic parts
                                if '[[italic_end]]' in part:
                                    italic_part, rest = part.split('[[italic_end]]', 1)
                                    # Add italic text
                                    italic_text = italic_part.encode('utf-8').decode('utf-8')
                                    italic_style = ParagraphStyle(
                                        name='GOSTItalic',
                                        parent=styles['GOSTNormal'],
                                        fontName='Times-Italic',
                                        encoding='UTF-8'
                                    )
                                    elements.append(Paragraph(italic_text, italic_style))
                                    # Add rest of the text
                                    if rest:
                                        text = rest.encode('utf-8').decode('utf-8')
                                        if text.strip().startswith('**') or text.strip().startswith('__'):
                                            text = text.strip('*_').strip()
                                            elements.append(Paragraph(text, styles['GOSTBold']))
                                        else:
                                            elements.append(Paragraph(text, styles['GOSTNormal']))
                                else:
                                    # If no end marker, treat as normal text
                                    text = part.encode('utf-8').decode('utf-8')
                                    if text.strip().startswith('**') or text.strip().startswith('__'):
                                        text = text.strip('*_').strip()
                                        elements.append(Paragraph(text, styles['GOSTBold']))
                                    else:
                                        elements.append(Paragraph(text, styles['GOSTNormal']))
                    else:
                        # Normal text without formatting
                        text = para.encode('utf-8').decode('utf-8')
                        if text.strip().startswith('**') or text.strip().startswith('__'):
                            text = text.strip('*_').strip()
                            elements.append(Paragraph(text, styles['GOSTBold']))
                        else:
                            elements.append(Paragraph(text, styles['GOSTNormal']))
                    elements.append(Spacer(1, 0))  # No extra space between paragraphs

    # Build the PDF
    pdf.build(elements)

    return output_path

