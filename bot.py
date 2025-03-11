import asyncio
import logging
import sys

from os import getenv, makedirs, path, remove
from docx import Document

from aiogram import Bot, Dispatcher, types, F
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from aiogram.filters import CommandStart
from aiogram.types import Message, FSInputFile
from dotenv import load_dotenv

from analys import analysis_conclusion, analysis_introduction, analysis_introduction_and_conclusion,analysis_literList
from CheckStructure import check_sections_in_docx, write_sections_to_docx


load_dotenv()
TOKEN = getenv('TOKEN')
DOWNLOAD_PATH = "files"
OUTPUT = "рекомендованные исправления.docx"


class Tg_Bot(Bot):

    # FIXME: тут один диспетчер для всех экземпляров. При создании в init(для каждого экземпляра) невозможно использовать декораторы
    dp = Dispatcher()
    section_keywords = ['список исполнителей', 'реферат',
                        'содержание', 'термины и определения', 'перечень сокращений и обозначений',
                        'введение', 'заключение',
                        'список использованных источников', 'приложения',
                        ]

    def __init__(self, output, path, token):
        """Дополняет init родительского класса, добавляя поля класса OUTPUT и DOWNLOAD_PATH"""
        super().__init__(token, default=DefaultBotProperties(parse_mode=ParseMode.HTML)) # вызов родительского init
        self.OUTPUT = output
        self.DOWNLOAD_PATH = path


    @dp.message(CommandStart())
    async def start_command(message: Message) -> None:
        """Приветствие, вызываемое командой /start"""
        await message.answer('Привет! Я бот который поможет тебе с твоей работой. '
                             'Отправь мне твою работу и я проверю её на соответствие ГОСТу')


    async def handle_document(self, message) -> None:
        """
        Сценарий обработки документа. Сохраняет документ на локальный диск,
        проводит его анализ и возвращает docx с рекомендациями.
        """
        file_id = message.document.file_id
        file_name = message.document.file_name

        makedirs(name=self.DOWNLOAD_PATH, exist_ok=True)

        file = await super().get_file(file_id)
        file_tg_path = file.file_path
        file_path = path.join(self.DOWNLOAD_PATH, file_name)

        await super().download_file(file_tg_path, file_path)
        await message.answer(f"Файл {file_name} успешно скачан! Начинаю анализ")

        logging.info(file_path)
        sections = check_sections_in_docx(file_path, self.section_keywords)
        write_sections_to_docx(self.OUTPUT, sections, file_path)
        new_doc = Document(self.OUTPUT)


        if sections.get('введение'):
            introduction = analysis_introduction(file_path, 'введение')
            new_doc.add_paragraph(introduction)
        if sections.get('заключение'):
            conclusion = analysis_conclusion(file_path, 'заключение')
            new_doc.add_paragraph(conclusion)
        if sections.get('введение') and sections.get('заключение'):
            introduction_and_conclusion = analysis_introduction_and_conclusion(file_path, 'введение', 'заключение')
            new_doc.add_paragraph(introduction_and_conclusion)
        if sections.get('список использованных источников'):
            liter_list = analysis_literList(file_path, 'список использованных источников')
            new_doc.add_paragraph(liter_list)

        new_doc.save(self.OUTPUT)
        doc = FSInputFile(self.OUTPUT)
        await message.answer_document(doc)

        remove(file_path) # удаляет файл пользователя после обработки

    def get_dp(self):
        """
        Возвращает диспетчера, созданного внутри класса
        ВОЗМОЖНО костыль
        """
        return self.dp



async def main() -> None:
    bot = Tg_Bot(token=TOKEN, path=DOWNLOAD_PATH, output=OUTPUT)
    dp = bot.get_dp()
    dp.message(F.document)(bot.handle_document) # используется здесь вместо декоратора, т.к. с декоратором невозможно использовать self
    await dp.start_polling(bot)

if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO, stream=sys.stdout)
    asyncio.run(main())