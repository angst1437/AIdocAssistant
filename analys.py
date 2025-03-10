from CheckStructure import get_text_after_section

from docx import Document

import ollama

def analysis_introduction(file_path, paragraphName):
    doc = Document(file_path)

    message = get_text_after_section(doc, paragraphName)
    message += 'Проверь текст на следующие требования. Если текст не соответствует какому-либо требованию '
    message += ' (или соответствует в неполном объёме, то предоставь информацию, '
    message += 'которую можно добавить в текст, чтобы он начал соответствовать требованию) '
    message += 'требования: оценка современного состояния решаемой научно-технической проблемы, '
    message += 'основание и исходные данные для разработки темы, обоснование необходимости проведения НИР, '
    message += 'сведения о планируемом научно-техническом уровне разработки, о патентных исследованиях и выводы из них, '
    message += 'актуальность и новизна темы, связь данной работы с другими научно-исследовательскими работами, '
    message += 'цель и задачи исследования'
    message += 'ответь на русском языке'

    try:
        formatted_message = [{"role": "user", "content": message}]
        response = ollama.generate(model='phi3:14b', messages=formatted_message)
        answer = response['message']['content']
        return 'Рекомендации по введению: \n' + answer

    except Exception as e:

        print('Произошла ошибка, попробуйте позже.')


def analysis_conclusion(file_path, paragraphName):
    doc = Document(file_path)

    message = get_text_after_section(doc, paragraphName)
    message += 'Проверь текст на следующие требования. Если текст не соответствует какому-либо требованию '
    message += ' (или соответствует в неполном объёме, то предоставь информацию, '
    message += 'которую можно добавить в текст, чтобы он начал соответствовать требованию) '
    message += 'требования: 1.краткие выводы по результатам работы или отдельных ее этапов,'
    message += '2. оценка полноты решений поставленных задач, 3. разработка рекомендаций и исходных данных по конкретному использованию результатов НИР'
    message += '4. результаты оценки технико-экономической эффективности внедрения'
    message += '5. результаты оценки научно-технического уровня выполненной НИР в сравнении с лучшими достижениями в этой области'
    message += 'ответь на русском языке'

    try:
        # Call the ollama.chat function with the context messages
        formatted_message = [{"role": "user", "content": message}]
        response = ollama.chat(model='phi3:14b', messages=formatted_message)
        answer = response['message']['content']
        return 'Рекомендации по заключению: \n'+answer

    except Exception as e:

        print('Произошла ошибка, попробуйте позже.')

def analysis_introduction_and_conclusion(file_path, paragraphName1, paragraphName2):
    doc = Document(file_path)

    message = 'введение: ' + get_text_after_section(doc, paragraphName1) + ' заключение: ' + get_text_after_section(doc, paragraphName2)
    message += 'проверь Пересечение Введения и Заключения. Цель и задачи выдвинутые во Введении должны соответствовать выводам сделанным в Заключении'
    message += ' если это не так, то предложи вариант исправления введения и заключения'
    message += 'ответ сделай на руссском языке'
    try:
        formatted_message = [{"role": "user", "content": message}]
        response = ollama.chat(model='phi3:14b', messages=formatted_message)
        answer = response['message']['content']
        return 'Рекомендации по связи введения и заключения: \n' + answer
    except Exception as e:

        print('Произошла ошибка, попробуйте позже.')
def analysis_literList(file_path, paragraphName):
    doc = Document(file_path)

    message = get_text_after_section(doc, 'список использованных источников')
    message += ' Проверь чтобы данные записи были оформлены по ГОСТ Р 7.0.100– 2018 '
    message += ' если какие-то записи не соответствуют оформлению, которые указаны в ГОСТ Р 7.0.100– 2018, то укажи их.'
    message += ' ответ дай на русском языке'

    try:
        formatted_message = [{"role": "user", "content": message}]
        response = ollama.chat(model='phi3:14b', messages=formatted_message)
        answer = response['message']['content']
        return ' Рекомендации по оформлению списка литературы: \n' + answer
    except Exception as e:

        print('Произошла ошибка, попробуйте позже.')