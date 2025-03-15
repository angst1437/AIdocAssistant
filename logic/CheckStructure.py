from docx import Document
from logic import CheckTitle

section_keywords = ['список исполнителей','реферат',
                    'содержание', 'термины и определения', 'перечень сокращений и обозначений',
                    'введение', 'заключение',
                    'список использованных источников','приложения',
                    ]

def get_headings(doc):
    headings = []  # Инициализируем список для заголовков
    for paragraph in doc.paragraphs:
        try:
            if paragraph.style.name.startswith('Heading'):
                headings.append(paragraph.text.lower())  # Добавляем только текст заголовка в список
        except AttributeError:
            pass
    return headings


def check_sections_in_docx(file_path, section_keywords):
    # Открываем документ
    print(file_path)
    try:
        doc = Document(file_path)
    except:
        doc = Document()
        doc.save(file_path)
    heads = get_headings(doc)

    # Словарь для хранения статуса разделов
    sections_found = {keyword: False for keyword in section_keywords}


    # Проходим по всем параграфам
    for paragraph in doc.paragraphs:
        is_bold = any(run.bold for run in paragraph.runs)


        # Проверяем, что параграф выровнен по центру и находится в начале страницы
        if is_bold:  # 1 соответствует центровке
            for keyword in section_keywords:
                if keyword.lower() == paragraph.text.lower() or keyword.lower() in heads :  # Сравниваем полное совпадение
                    sections_found[keyword] = True

    return sections_found


def get_text_after_section(doc, section_name):
    text_after_section = ""
    is_in_section = False  # Флаг, указывающий, что мы находимся в целевой секции
    heads = get_headings(doc)

    sections_found = {keyword: False for keyword in section_keywords}

    # Проходим по всем параграфам в документе
    for paragraph in doc.paragraphs:
        # Если мы уже в секции, продолжаем собирать параграфы
        if is_in_section:
            if paragraph.text.lower() in section_keywords or paragraph.text.lower() in heads:
                break
            # Добавляем текст параграфа в строку
            text_after_section += paragraph.text

        # Проверяем, если текущий параграф — это заголовок нужной секции
        if paragraph.text.lower() == section_name.lower():
            is_in_section = True  # Начинаем собирать текст после этого заголовка

    return text_after_section.strip()  # Убираем лишние пробелы в начале и конце строки

def write_sections_to_docx(output_file, sections_status, file_path):
    # Создаем новый документ
    words = CheckTitle.extract_words_from_first_page(file_path)

    # Проверяем наличие титульного листа
    title_page_exists = CheckTitle.check_title_page(words)

    # Создаем новый документ
    doc = Document()

    # Добавляем строку с результатом
    result_sign = '+' if title_page_exists else '-'
    doc.add_paragraph(f"Титульный лист: {result_sign}")

    # Для каждого ключевого слова создаем строку и записываем в документ
    for keyword, status in sections_status.items():
        sign = '+' if status else '-'
        doc.add_paragraph(f"{keyword}: {sign}")

    # Сохраняем документ
    doc.add_paragraph('Разделы, отмеченные минусом либо отсутствуют в вашей работе, либо некорректен заголовок раздела (возможно он отсутствует или он не выделен полужирным или ещё что-то некорректно с заголовком)')
    doc.save(output_file)
