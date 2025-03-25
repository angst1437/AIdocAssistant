from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.shared import Pt
import logic.datatypes as datatypes


def mm_to_inches(mm):
    return Inches(mm / 25.4)

class DocBuilder:
    def __init__(self, parts=None):
        self.doc = Document()
        self.parts = parts
        self.gost_style = self.get_gost_style()


    def get_doc(self):
        return self.doc

    def build(self, blocks: dict):
        self.title_builder(blocks['title'])
        self.content_builder(blocks['content'])
        self.intro_builder(blocks['intro'])
        self.change_indent()
        self.doc.save("test.docx")

    def change_indent(self):
        for section in self.doc.sections:
            section.left_margin = mm_to_inches(30)  # Левое поле 30 мм
            section.right_margin = mm_to_inches(15)  # Правое поле 15 мм
            section.top_margin = mm_to_inches(20)  # Верхнее поле 20 мм
            section.bottom_margin = mm_to_inches(20)  # Нижнее поле 20 мм

    def get_gost_style(self):
        gost_style = self.doc.styles.add_style('GOST', WD_STYLE_TYPE.PARAGRAPH)
        gost_style.font.name = 'Times New Roman'
        gost_style.font.size = Pt(14)
        gost_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        gost_style.paragraph_format.first_line_indent = mm_to_inches(12.5)
        return gost_style

    def make_heading(self, text, is_bold=True, is_centered=True):
        p = self.doc.add_paragraph(style=self.gost_style)
        run = p.add_run(text)
        run.bold = is_bold
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if is_centered else WD_PARAGRAPH_ALIGNMENT.RIGHT


    def title_builder(self, title_data: datatypes.TitleData or None=None):
        if self.parts is not None:
            pass
        else:
            self.make_heading("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РФ", is_bold=False)

            for val in title_data.university_info:
                self.make_heading(val, is_bold=False)

            self.make_heading(("\n" * 5) + title_data.work_type.upper())
            self.make_heading("По дисциплине " + title_data.subject.strip())
            self.make_heading("Тема: " + title_data.theme)
            self.make_heading(("\n" * 4) + "Выполнена студентом:\n" + title_data.author, is_centered=False, is_bold=False)
            self.make_heading("Группа: " + title_data.group, is_centered=False, is_bold=False)
            self.make_heading("Руководитель: " + title_data.educator, is_centered=False, is_bold=False)
            self.make_heading(("\n" * 2) + title_data.city, is_bold=False)
            self.make_heading(title_data.year, is_bold=False)


    def content_builder(self, text):
        self.make_heading("СОДЕРЖАНИЕ")
        for p in text:
            self.doc.add_paragraph(p, style=self.gost_style)


    def intro_builder(self, text):
        self.make_heading("Введение")
        for p in text:
            self.doc.add_paragraph(p, style=self.gost_style)


    def main_part_builder(self):
        pass

    def conclusion_builder(self):
        pass

    def sources_builder(self):
        pass